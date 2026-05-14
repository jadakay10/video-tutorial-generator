import json
import re
import threading
import queue
import uuid
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt
from flask import Flask, request, jsonify, Response, render_template, send_file
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

from video_to_tutorial import (
    extract_audio,
    extract_frames,
    sample_frames,
    transcribe_audio,
    encode_frames,
    generate_tutorial,
)

load_dotenv()

app = Flask(__name__)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# job_id -> Queue of SSE event dicts
_jobs: dict[str, queue.Queue] = {}
_active_job: list[str] = []  # at most one entry
_lock = threading.Lock()


def _markdown_to_docx(markdown_text: str) -> bytes:
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.strip() in ("", "---"):
            continue
        else:
            p = doc.add_paragraph()
            for i, part in enumerate(re.split(r"\*\*(.*?)\*\*", line)):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _run_pipeline(job_id: str, video_path: Path) -> None:
    q = _jobs[job_id]
    try:
        audio_path = video_path.parent / (video_path.stem + ".mp3")
        frames_dir = video_path.parent / (video_path.stem + "_frames")

        q.put({"step": 1, "label": "Extracting audio"})
        extract_audio(video_path, audio_path)

        q.put({"step": 2, "label": "Extracting frames"})
        frame_files = extract_frames(video_path, frames_dir)

        q.put({"step": 3, "label": "Transcribing with Whisper"})
        transcript = transcribe_audio(audio_path)

        q.put({"step": 4, "label": "Encoding frames"})
        frame_files = sample_frames(frame_files)
        image_blocks = encode_frames(frame_files)

        q.put({"step": 5, "label": "Generating tutorial with Claude"})
        tutorial = generate_tutorial(transcript, image_blocks)

        markdown = f"## Transcript\n\n{transcript}\n\n---\n\n{tutorial}"
        q.put({
            "done": True,
            "markdown": markdown,
            "transcript": transcript,
            "tutorial": tutorial,
            "filename": video_path.stem,
        })
    except Exception as exc:
        q.put({"error": str(exc)})
    finally:
        if _active_job and _active_job[0] == job_id:
            _active_job.clear()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    if not file or file.filename == "":
        return jsonify({"error": "No file provided."}), 400
    if not file.filename.lower().endswith(".mp4"):
        return jsonify({"error": "Only .mp4 files are supported."}), 400

    with _lock:
        if _active_job:
            return jsonify({"error": "A job is already running. Please wait."}), 409

        job_id = str(uuid.uuid4())
        job_dir = UPLOAD_DIR / job_id
        job_dir.mkdir(parents=True)
        video_path = job_dir / secure_filename(file.filename)
        file.save(video_path)

        _jobs[job_id] = queue.Queue()
        _active_job.append(job_id)

    try:
        thread = threading.Thread(target=_run_pipeline, args=(job_id, video_path), daemon=True)
        thread.start()
    except Exception:
        with _lock:
            if _active_job and _active_job[0] == job_id:
                _active_job.clear()
        return jsonify({"error": "Failed to start processing thread."}), 500

    return jsonify({"job_id": job_id})


@app.route("/progress/<job_id>")
def progress(job_id: str):
    if job_id not in _jobs:
        return jsonify({"error": "Job not found."}), 404

    def generate():
        q = _jobs[job_id]
        while True:
            try:
                event = q.get(timeout=300)
            except queue.Empty:
                yield f"data: {json.dumps({'error': 'Processing timed out.'})}\n\n"
                break
            yield f"data: {json.dumps(event)}\n\n"
            if "done" in event or "error" in event:
                break

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/download/docx", methods=["POST"])
def download_docx():
    data = request.get_json()
    if not data or "markdown" not in data:
        return jsonify({"error": "No markdown provided."}), 400
    filename = secure_filename(data.get("filename", "tutorial")) + ".docx"
    docx_bytes = _markdown_to_docx(data["markdown"])
    return send_file(
        BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


if __name__ == "__main__":
    app.run(debug=True, port=5000)
